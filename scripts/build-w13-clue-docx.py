# -*- coding: utf-8 -*-
"""Generate docs/w13-clue-guide.docx from guide content."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "W13线索对应攻略.docx"


def set_doc_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_title(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_quote(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Pt(18)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.italic = True


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.left_indent = Pt(12)
        for run in p.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(10)


def main() -> None:
    doc = Document()
    set_doc_font(doc)

    add_title(doc, "W13 线索对应攻略")
    add_quote(
        doc,
        "《不存在的学生》提交页（/audit/submit · W13）8 条核心线索 + 2 条附加线索与游戏内页面的完整对应关系。\n"
        "面向：自行推理的玩家、测试同学、课程答辩演示。",
    )
    doc.add_paragraph()

    # 一、总览
    add_heading(doc, "一、总览")
    add_table(
        doc,
        ["类型", "数量", "是否提交必需"],
        [
            ["核心线索", "8 条", "是（8/8 才能选「更正为创作路径账户」）"],
            ["附加线索", "2 条", "否（真结局二需 open_path_code）"],
        ],
    )
    add_table(
        doc,
        ["分组", "线索数", "审核终端对应核验线"],
        [
            ["身份核验", "2", "学籍身份"],
            ["项目溯源", "2", "项目成果"],
            ["AI 与创作路径", "4", "AI 与创作路径 + 账号性质"],
        ],
    )

    # 二、身份核验
    add_heading(doc, "二、身份核验（2 条）")
    add_table(
        doc,
        ["#", "线索内容", "对应页面", "怎么触发"],
        [
            [
                "1",
                "程野无任何真实学籍记录",
                "审核工作台 · W04 学生档案查询 /audit/student",
                "输入学号 2034020103，打开程野档案",
            ],
            [
                "2",
                "内部编号 CYA-0000 并非学生账号",
                "同上",
                "同一档案页，可见内部编号 CYA-0000",
            ],
        ],
    )
    add_para(doc, "推荐入口", bold=True)
    add_para(doc, "1. /audit/graduates（W03 优秀毕业生名单）找到程野学号")
    add_para(doc, "2. → /audit/student 查询 2034020103")

    # 三、项目溯源
    add_heading(doc, "三、项目溯源（2 条）")
    add_table(
        doc,
        ["#", "线索内容", "对应页面", "怎么触发"],
        [
            [
                "3",
                "获奖新闻中姓名与照片出现异常",
                "学院官网 · 创新成果奖新闻 /anima/news/innovation-award-2034",
                "打开新闻，点击/悬停团队表格触发 glitch 异常",
            ],
            [
                "4",
                "真实创作者为林澈及跨专业团队",
                "多处均可（见下表）",
                "在以下任一页面触发即可",
            ],
        ],
    )
    add_para(doc, "project_lin_che 可触发的全部页面", bold=True)
    add_table(
        doc,
        ["页面", "路径", "操作"],
        [
            ["创新成果奖新闻", "/anima/news/innovation-award-2034", "阅读新闻正文"],
            ["山海行灯毕业新闻", "/anima/news/shanhai-lamp-graduation-2034", "阅读新闻正文"],
            ["学生档案查询", "/audit/student", "查询学号 2030120401（林澈）"],
            ["实验室申请记录", "/audit/lab", "查询学号 2030120401"],
            ["学生项目查询", "/audit/project", "查询 CYA-P-2034-01（《未命名之路》立项书）"],
        ],
    )
    add_para(doc, "推荐入口", bold=True)
    add_para(doc, "· 学院官网搜索「未命名之路」或「创新成果」→ 打开创新成果奖新闻")
    add_para(doc, "· 或 /audit/project 直接查 CYA-P-2034-01")

    # 四、AI 与创作路径
    add_heading(doc, "四、AI 与创作路径（4 条）")
    add_table(
        doc,
        ["#", "线索内容", "对应页面", "怎么触发"],
        [
            [
                "5",
                "Ani AI 负责原型测试与资源匹配",
                "审核工作台 · W12 学生项目查询 /audit/project",
                "查询 CYA-P-2034-02，打开《山海行灯》立项书",
            ],
            [
                "6",
                "林澈与韩老师的跨媒介聊天记录",
                "W12 → 内部附件 /audit/attachments",
                "山海行灯立项书底部「查看内部附件 →」",
            ],
            [
                "7",
                "「创作路径账号」关键词",
                "同上 /audit/attachments",
                "附件页聊天记录中出现该词",
            ],
            [
                "8",
                "程野是误识别的创作路径账号",
                "W12 /audit/project",
                "查询 CYA-P-2030-01 或 CYA-0000，打开 CYA 主项目立项书",
            ],
        ],
    )
    add_para(doc, "推荐调查顺序", bold=True)
    add_code_block(
        doc,
        [
            "/audit/project  →  CYA-P-2034-02（山海行灯）",
            "       ↓",
            "/audit/attachments（内部附件 · 聊天证据）",
            "       ↓",
            "/audit/project  →  CYA-P-2030-01 或 CYA-0000（CYA 真相）",
        ],
    )
    add_para(doc, "补充说明", bold=True)
    add_para(
        doc,
        "· 学院官网 /anima/news/ani-ai-launch 与 /anima/education 也能触发部分 AI 相关线索，"
        "但审核线主路径以 W12 + 内部附件为准。",
    )
    add_para(
        doc,
        "· 登录程野视角学生系统 /student/home（账号 CYA-0000）也会触发 cya_truth，"
        "与 W12 查 CYA 立项书等效。",
    )

    # 五、附加线索
    add_heading(doc, "五、附加线索（2 条 · 非提交必需）")
    add_table(
        doc,
        ["#", "线索内容", "对应页面", "怎么触发", "用途"],
        [
            [
                "9",
                "个人毕业项目多次退回",
                "学生个人系统 · 毕业项目管理 /student/project",
                "祁玉账号登录（学号 203508083038），非程野",
                "氛围 / 侧写",
            ],
            [
                "10",
                "需向 Ani AI 询问项目招募码",
                "学生个人系统 · Ani AI /student/home?chat=ani",
                "打开 Ani 对话，询问招募码 / 开放创作路径",
                "真结局二必需",
            ],
        ],
    )
    add_para(doc, "真结局二额外条件", bold=True)
    add_para(doc, "· 向 Ani 问到招募码 PATH-2034-LC 后自动获得 open_path_code")
    add_para(
        doc,
        "· 提交「更正为创作路径账户」→ 反思第 5 题选认可 → /audit/path-invite → /audit/next-path",
    )

    # 六、推荐完整路线
    add_heading(doc, "六、推荐完整路线（速查）")
    add_para(doc, "适合第一次通关、按顺序收齐 8 条核心线索：")
    add_code_block(
        doc,
        [
            "① 邮箱开局 → 登录审核工作台",
            "        ↓",
            "② W03 /audit/graduates",
            "   发现程野（2034020103）→ 操作「信息有误」（解锁审核终端）",
            "        ↓",
            "③ W04 /audit/student",
            "   查 2034020103 → 身份 2 条线索",
            "        ↓",
            "④ 学院官网 /anima/search",
            "   搜「未命名之路」→ /anima/news/innovation-award-2034",
            "   → 项目异常 + 林澈（可顺带读山海行灯新闻）",
            "        ↓",
            "⑤ W12 /audit/project",
            "   查 CYA-P-2034-02 → 山海行灯立项书",
            "        ↓",
            "⑥ /audit/attachments",
            "   内部附件 → AI 聊天 2 条 + 创作路径账号关键词",
            "        ↓",
            "⑦ W12 /audit/project",
            "   查 CYA-P-2030-01 或 CYA-0000 → CYA 真相",
            "        ↓",
            "⑧ W13 /audit/submit",
            "   8/8 核心线索 → 「更正为创作路径账户」",
            "        ↓",
            "⑨ /audit/reflection → 真结局 / 坏结局分叉",
        ],
    )
    add_para(doc, "（可选 · 真结局二）步骤 ⑧ 之前：学生系统 Ani 对话拿到 PATH-2034-LC。")

    # 七、关键编号速查
    add_heading(doc, "七、关键编号速查")
    add_heading(doc, "学号", level=2)
    add_table(
        doc,
        ["人物", "学号", "用途"],
        [
            ["程野", "2034020103", "W03 名单、W04 学籍异常"],
            ["林澈", "2030120401", "W04 档案、W07 实验室"],
            ["祁玉（玩家）", "203508083038", "学生系统、Ani 对话"],
        ],
    )
    add_heading(doc, "项目编码", level=2)
    add_table(
        doc,
        ["编码", "名称", "主要线索"],
        [
            ["CYA-P-2034-01", "《未命名之路》", "林澈与团队"],
            ["CYA-P-2034-02", "《山海行灯》", "Ani AI → 内部附件"],
            ["CYA-P-2030-01", "Creative Yard 主项目", "CYA 真相"],
            ["CYA-0000", "创作路径账号", "同上（自动映射至 CYA-P-2030-01）"],
        ],
    )
    add_heading(doc, "账号", level=2)
    add_table(
        doc,
        ["系统", "账号", "密码"],
        [
            ["审核工作台", "DMXY2036QD", "CMU@DH2036"],
            ["学生系统（祁玉）", "203508083038", "Project2024!"],
            ["创作路径（程野）", "CYA-0000", "CreativeYard2030"],
        ],
    )

    # 八、核验线
    add_heading(doc, "八、与审核终端四条核验线的对应")
    add_table(
        doc,
        ["核验线", "需要的线索（任一条点亮该线；账号线除外）", "账号线特殊要求"],
        [
            ["学籍身份", "identity_no_enrollment 或 identity_cya_code", "—"],
            ["项目成果", "project_award_anomaly 或 project_lin_che", "—"],
            ["AI 与创作路径", "ai_ani_platform / ai_cross_media / ai_creative_path 任一", "—"],
            ["账号性质", "—", "必须 cya_truth"],
        ],
    )
    add_quote(doc, "在 W03 对程野点「信息有误」才会启动审核终端；仅收线索而不标记异常，无法走真结局提交线。")

    # 九、结局
    add_heading(doc, "九、结局与线索关系（简表）")
    add_table(
        doc,
        ["结局", "名称", "核心线索", "终端", "提交选项", "反思 Q5", "招募码"],
        [
            ["坏 A", "《无事发生的暑假》", "任意", "未解锁", "信息无误", "—", "—"],
            ["坏 B", "《不存在的人》", "8/8", "已解锁", "删除申请 或 Q5 不认可", "—", "—"],
            ["真 1", "《被看见的路径》", "8/8", "已解锁", "更正为路径账户", "认可", "无"],
            ["真 2", "《下一段路径》", "8/8", "已解锁", "更正为路径账户", "认可", "有 PATH-2034-LC"],
        ],
    )
    add_para(doc, "更完整的结局流程见 docs/player-walkthrough.md")

    # 页脚
    doc.add_paragraph()
    p = doc.add_paragraph("文档版本与当前构建一致。若后续调整线索或路由，以游戏内实际表现为准。")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.font.size = Pt(9)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Written: {OUT}")


if __name__ == "__main__":
    main()
