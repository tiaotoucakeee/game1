# -*- coding: utf-8 -*-
"""导出「你好！祁玉」开发者邮件为 docx。"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

DEV_GROUP = "《动画学院暑期勤工俭学事件》开发组"
OUTPUT = Path(__file__).resolve().parent.parent / "docs" / "你好！祁玉-开发者邮件.docx"

META = {
    "subject": "你好！祁玉",
    "from": DEV_GROUP,
    "from_email": "dev@workstudy-event.game",
    "to": "qiyu_3608@cuc.edu.cn",
    "date": "？？？",
}

BODY = [
    ("lead", "祁玉，你好。"),
    (
        "lead",
        "如果你正在读这封信，说明你已经进入了这场平行实境调查——一款 ARG 网页解密游戏。下面是我们为你准备的引言与玩法说明，建议在开始审核任务前先读一遍。",
    ),
    ("blank", ""),
    ("heading", "游戏引言"),
    (
        "normal",
        "在生成式人工智能快速发展的背景下，动画、漫画、数字媒体、游戏与交互艺术的边界逐渐模糊，传统专业的调整与合并也使部分艺术类专业学生产生「原有专业是否失去价值」「AI 是否会替代创作者」等焦虑。",
    ),
    (
        "normal",
        "游戏以中国传媒大学动画与数字艺术学院的未来官网为虚构背景，通过一场「不存在的学生档案」事件，引导玩家重新理解专业调整、媒介融合与 AI 辅助创作的关系。",
    ),
    (
        "emphasis",
        "你将扮演学院数字媒体中心的临时审核员，负责核验即将发布的优秀毕业生名单。审核中，你会发现名为「程野」的学生拥有完整的四年课程、项目与获奖记录，却没有任何录取和注册信息。你需要登录管理员后台，通过搜索编号、查看历史页面、核对学籍档案与专业文件，逐步还原其真实身份。",
    ),
    (
        "normal",
        "游戏强调 AI 作为创作工具的积极价值：它能够帮助学生验证创意、转换媒介、匹配资源并降低跨领域实践门槛，但作品的主题、审美与最终方向仍由创作者决定。",
    ),
    (
        "emphasis",
        "我们希望将玩家对「专业取消」和「AI 发展」的焦虑，转化为对未来艺术教育的积极理解——被取消的不是各专业的价值，而是它们之间难以跨越的边界。AI 拓展创作的可能，而创作者决定这些可能通向哪里。",
    ),
    ("blank", ""),
    ("heading", "玩法介绍"),
    (
        "normal",
        "你的任务是核查优秀毕业生名单中异常出现的「程野」。游戏以学院官网为主要界面，你可以通过搜索姓名、学号、项目编号和文件关键词，进入学籍系统、历史新闻、项目档案等不同页面，在浏览过程中发现并记录线索。",
    ),
    (
        "emphasis",
        "登录审核工作台后，可在右下角找到「审核终端」。它用于提示学籍、项目归属和账号性质等信息是否核验完成，但不会直接显示游戏进度条。",
    ),
    (
        "emphasis",
        "调查大致分为三条线索线，最终在 CYA-0000 账号处汇合：",
    ),
    (
        "normal",
        "1. 学籍身份线：程野档案 → 学号与专业异常 → 学籍核验 → 无录取、注册、班级、校园卡记录 → CYA-0000 并非学生账号。结论：程野不是一名真实学生。",
    ),
    (
        "normal",
        "2. 项目成果线：程野档案中的获奖名称 → 搜索奖项 → 进入学院新闻 → 获奖名单和照片出现异常 → 搜索林澈 → 找到实验室申请记录。结论：这些作品由林澈和不同专业的真实学生共同完成，并非程野个人成果。",
    ),
    (
        "normal",
        "3. AI 与创作路径线：搜索 Ani AI → 进入智能畅课平台 → 查看优秀案例 → 搜索「跨媒介叙事」→ 找到林澈项目与韩老师聊天记录 → 获得「创作路径账号」与 CYA-0000。结论：AI 只负责原型测试、课程匹配与技术辅助，真正作出创作决定的是学生。",
    ),
    (
        "normal",
        "你可以随时提交审核结论。根据已收集的关键线索和你的判断，游戏会走向不同结局：误认程野为 AI 伪造学生并删除档案、查明程野是跨专业创作路径账号，或进一步登录个人学生系统、与 Ani 智能体对话，找到隐藏入口并加入项目组。",
    ),
    (
        "emphasis",
        "整体玩法以自由搜索和网页调查为主，不是线性点击解谜。建议先从录用通知中的审核账号进入工作台，再前往学院官网开始搜索。",
    ),
    ("blank", ""),
    ("lead", "祝调查顺利。"),
    ("blank", ""),
    ("emphasis", f"—— {DEV_GROUP}"),
]


def set_run_font(run, size_pt=12, bold=False, east_asia="宋体"):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def add_meta_line(doc, label, value):
    p = doc.add_paragraph()
    label_run = p.add_run(f"{label}：")
    label_run.bold = True
    set_run_font(label_run, 11)
    value_run = p.add_run(value)
    set_run_font(value_run, 11)
    p.paragraph_format.space_after = Pt(4)


def add_body_paragraph(doc, tone, text):
    if tone == "blank":
        doc.add_paragraph()
        return

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(8)

    if tone == "heading":
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run(text)
        set_run_font(run, 14, bold=True, east_asia="黑体")
        return

    size = 12 if tone == "normal" else 12
    bold = tone in ("emphasis", "lead")
    run = p.add_run(text)
    set_run_font(run, size, bold=bold)


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    title = doc.add_paragraph()
    title.alignment = 0
    title_run = title.add_run(META["subject"])
    set_run_font(title_run, 18, bold=True, east_asia="黑体")
    title.paragraph_format.space_after = Pt(16)

    add_meta_line(doc, "发件人", f'{META["from"]} <{META["from_email"]}>')
    add_meta_line(doc, "收件人", META["to"])
    add_meta_line(doc, "时间", META["date"])

    doc.add_paragraph()

    for tone, text in BODY:
        add_body_paragraph(doc, tone, text)

    doc.save(OUTPUT)
    print(f"已保存：{OUTPUT}")


if __name__ == "__main__":
    main()
